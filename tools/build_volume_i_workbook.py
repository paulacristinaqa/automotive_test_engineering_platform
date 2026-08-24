from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "workbook-volume-i.md"
OUTPUT = ROOT / "docs" / "ATEP_Volume_I_Core_Platform_Engineering_Workbook.docx"
VOLUME_NUMBER = "I"
VOLUME_NAME = "Core Platform"
DOCUMENT_VERSION = "1.0.0"
DOCUMENT_STATUS = "Living document - Volume I baseline implemented"
BASELINE_DATE = "12 August 2026"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C4D1"
WHITE = "FFFFFF"
BLACK = "111827"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(
    run,
    size: float,
    color: str = BLACK,
    bold: bool | None = None,
    italic: bool | None = None,
    name: str = "Calibri",
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_font(run, 8.5, MUTED)


def create_abstract_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    existing = [
        int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022" if bullet else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstract)
    else:
        numbering.append(abstract)
    return abstract_id


def new_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num = numbering.add_num(abstract_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    return int(num.numId)


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    level = num_pr.get_or_add_ilvl()
    level.val = 0
    identifier = num_pr.get_or_add_numId()
    identifier.val = num_id


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"ATEP  /  VOLUME {VOLUME_NUMBER}  /  {VOLUME_NAME.upper()}")
        set_font(run, 8.5, MUTED, bold=True)

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.paragraph_format.space_after = Pt(0)
        footer_paragraph.paragraph_format.tab_stops.clear_all()
        footer_paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
        )
        set_font(
            footer_paragraph.add_run(
                f"ATEP Volume {VOLUME_NUMBER} Workbook  |  Version {DOCUMENT_VERSION}"
            ),
            8.5,
            MUTED,
        )
        # LibreOffice preserves the Footer style's center/right tab stops even when
        # direct paragraph tabs are present. Two tabs reliably reach the right stop.
        set_font(footer_paragraph.add_run("\t\tPage "), 8.5, MUTED)
        add_page_field(footer_paragraph)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_font(kicker.add_run("AUTOMOTIVE TEST ENGINEERING PLATFORM"), 10, BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run(f"Volume {VOLUME_NUMBER}"), 30, NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    set_font(subtitle.add_run(f"{VOLUME_NAME} Engineering Workbook"), 17, DARK_BLUE, bold=True)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(34)
    set_font(
        descriptor.add_run(
            "Architecture, implementation record, verification strategy, and engineering evidence"
        ),
        11.5,
        MUTED,
        italic=True,
    )

    meta = doc.add_table(rows=4, cols=2)
    rows = [
        ("Document version", DOCUMENT_VERSION),
        ("Baseline date", BASELINE_DATE),
        ("Status", DOCUMENT_STATUS),
        ("Language", "English"),
    ]
    for row, (label, value) in zip(meta.rows, rows, strict=True):
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(meta, [2700, 6660], compact=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        p.add_run(
            "A living technical record for architecture, software engineering, "
            "testing, operations, and portfolio evidence."
        ),
        10.5,
        MUTED,
    )
    doc.add_page_break()


def add_contents(doc: Document, lines: list[str], decimal_abstract_id: int) -> None:
    doc.add_heading("Contents", level=1)
    intro = doc.add_paragraph(
        "This contents list follows the workbook's maintained section structure. "
        "Word's Navigation Pane can also be used because all sections use semantic "
        "heading styles."
    )
    intro.paragraph_format.space_after = Pt(10)
    contents_num_id = new_numbering_instance(doc, decimal_abstract_id)
    for line in lines:
        if line.startswith("## "):
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, contents_num_id)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.first_line_indent = Inches(0)
            text = line[3:].strip()
            set_font(p.add_run(text), 10.5, NAVY, bold=True)
    doc.add_page_break()


def add_inline_markdown(paragraph, text: str, size: float | None = None) -> None:
    text = text.replace("—", "-").replace("–", "-").replace("→", "->")
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size or 11, BLACK, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size or 9.5, DARK_BLUE, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_font(run, size or 11, BLACK)


def column_widths(column_count: int) -> list[int]:
    patterns = {
        2: [2700, 6660],
        3: [1650, 4450, 3260],
        4: [1150, 3700, 3100, 1410],
        5: [850, 2600, 2200, 2650, 1060],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def style_table(table, widths: list[int], compact: bool = True) -> None:
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    size = 8.1 if compact and len(widths) >= 4 else 8.8 if compact else 9.5
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, LIGHT_BLUE if row_index == 0 else WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if cell_index == 0 and len(widths) >= 3
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_font(run, size, NAVY if row_index == 0 else BLACK, bold=row_index == 0)
    after = table._tbl.getnext()
    if after is None or after.tag != qn("w:p"):
        table._parent.add_paragraph().paragraph_format.space_after = Pt(3)


def select_column_widths(rows: list[list[str]]) -> list[int]:
    header = rows[0]
    if header[0] == "Method and route":
        return [3100, 4000, 2260]
    if header[0] == "Variable":
        return [2400, 1500, 1500, 3960]
    if header[0] == "Technology":
        return [1300, 2400, 3100, 2560]
    if header[0] == "Version":
        return [1150, 1400, 3450, 3360]
    if header[0] == "ID" and len(header) == 4:
        if header[1] == "Requirement":
            return [1150, 3550, 3100, 1560]
        return [1150, 3700, 3100, 1410]
    return column_widths(len(header))


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            paragraph = table.cell(row_index, col_index).paragraphs[0]
            paragraph.clear()
            add_inline_markdown(paragraph, value, size=8.2 if len(row) >= 4 else 9)
    style_table(table, select_column_widths(rows))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1
    rows: list[list[str]] = []
    for line in raw:
        values = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    return rows, index


def render_body(
    doc: Document, lines: list[str], decimal_abstract_id: int, bullet_abstract_id: int
) -> None:
    index = next(i for i, line in enumerate(lines) if line.startswith("## "))
    page_break_sections = {
        "8. Requirements and Traceability",
        "10. Test Strategy",
        "11. Detailed Test Catalogue",
        "11.1 III-2 Verification",
        "12. Implemented Evidence",
        "12. Suggested CI/CD Quality Pipeline",
        "16. Engineering Review Worksheets",
        "18. Glossary",
    }
    number_num_id: int | None = None
    bullet_num_id: int | None = None
    previous_kind = ""
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            number_num_id = None
            bullet_num_id = None
            text = line[3:].strip()
            heading = doc.add_heading(text, level=1)
            if text in page_break_sections and len(doc.paragraphs) > 2:
                heading.paragraph_format.page_break_before = True
            index += 1
            continue
        if line.startswith("### "):
            number_num_id = None
            bullet_num_id = None
            doc.add_heading(line[4:].strip(), level=2)
            index += 1
            continue
        if line.startswith("#### "):
            number_num_id = None
            bullet_num_id = None
            doc.add_heading(line[5:].strip(), level=3)
            index += 1
            continue
        if line.startswith("|"):
            number_num_id = None
            bullet_num_id = None
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        if line.startswith("- "):
            number_num_id = None
            if previous_kind != "bullet" or bullet_num_id is None:
                bullet_num_id = new_numbering_instance(doc, bullet_abstract_id)
            p = doc.add_paragraph(style="List Bullet")
            apply_numbering(p, bullet_num_id)
            add_inline_markdown(p, line[2:].strip())
            previous_kind = "bullet"
            index += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            bullet_num_id = None
            if previous_kind != "number" or number_num_id is None:
                number_num_id = new_numbering_instance(doc, decimal_abstract_id)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, number_num_id)
            add_inline_markdown(p, number_match.group(1))
            previous_kind = "number"
            index += 1
            continue

        number_num_id = None
        bullet_num_id = None
        previous_kind = "paragraph"
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith(("#", "|", "- "))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, " ".join(paragraph_lines))


def build() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    properties = doc.core_properties
    properties.title = f"ATEP Volume {VOLUME_NUMBER} - {VOLUME_NAME} Engineering Workbook"
    properties.subject = "Architecture, implementation, verification, and engineering evidence"
    properties.author = "ATEP Core Platform Engineering"
    properties.keywords = (
        "ATEP, automotive testing, FastAPI, RBAC, PostgreSQL, RabbitMQ, engineering workbook"
    )
    properties.comments = f"Generated from {SOURCE.relative_to(ROOT).as_posix()}"

    add_cover(doc)
    decimal_abstract_id = create_abstract_numbering(doc, bullet=False)
    bullet_abstract_id = create_abstract_numbering(doc, bullet=True)
    add_contents(doc, lines, decimal_abstract_id)
    render_body(doc, lines, decimal_abstract_id, bullet_abstract_id)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if paragraph.text.startswith("Important limitation:"):
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.18)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
